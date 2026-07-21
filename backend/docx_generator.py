"""
docx_generator.py

Generates a Microsoft Word (.docx) question paper.
"""

import os
from docx import Document
from docx.shared import Pt
from config import Config


class DOCXGenerator:

    @staticmethod
    def generate(
        question_paper,
        school_name,
        class_name,
        subject,
        exam_name,
        duration,
        max_marks,
        filename="Question_Paper.docx"
    ):

        os.makedirs(Config.GENERATED_FOLDER, exist_ok=True)

        filepath = os.path.join(
            Config.GENERATED_FOLDER,
            filename
        )

        document = Document()

        # -------------------------
        # Title
        # -------------------------

        title = document.add_heading(school_name, level=1)
        title.alignment = 1

        heading = document.add_heading(exam_name, level=2)
        heading.alignment = 1

        p = document.add_paragraph()

        p.add_run("Class : ").bold = True
        p.add_run(class_name)

        p.add_run("        Subject : ").bold = True
        p.add_run(subject)

        p = document.add_paragraph()

        p.add_run("Time : ").bold = True
        p.add_run(duration)

        p.add_run("        Maximum Marks : ").bold = True
        p.add_run(str(max_marks))

        document.add_paragraph()

        current_lesson = None

        # -------------------------
        # Questions
        # -------------------------

        for q in question_paper:

            if current_lesson != q["lesson"]:

                current_lesson = q["lesson"]

                document.add_heading(current_lesson, level=3)

            para = document.add_paragraph()

            para.style.font.size = Pt(12)

            para.add_run(
                f'{q["question_no"]}. '
            ).bold = True

            para.add_run(q["question"])

            para.add_run(
                f'    [{q["marks"]} Marks]'
            ).bold = True

        document.save(filepath)

        return filepath
